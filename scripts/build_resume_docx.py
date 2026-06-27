#!/usr/bin/env python3
"""Generate Praisel Ekpenyong SOC resume (.docx) — Headless Headhunter template."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Praisel_Ekpenyong_Resume.docx"
OUTPUT_FALLBACK = ROOT / "docs" / "Praisel_Ekpenyong_Resume_v2.docx"

FONT = "Arial"
BODY_SIZE = Pt(10.5)
NAME_SIZE = Pt(14)
CONTACT_SIZE = Pt(12)
LINE_SPACING = 1.5

CONTACT = {
    "name": "Praisel Richard Ekpenyong",
    "phone": "647-472-9928",
    "email": "ekpenyongpraisel@gmail.com",
    "links": "linkedin.com/in/praiselekpenyong | praisel-ekpenyong.github.io/Portfolio",
    "location": "Authorized to work in Canada at Edmonton, Alberta",
}

EDUCATION = [
    "B.Sc. in Cybersecurity from Miva University, Abuja Status - Graduated 2025",
    "CompTIA Security+, Microsoft SC-200, TryHackMe SOC Level 1",
]

PROJECTS = [
    {
        "title": "Independent On-Premises SOC Homelab",
        "dates": None,
        "bullets": [
            "Monitored Windows and Linux endpoints with Splunk, Wazuh, and Active Directory, performed alert triage through security monitoring and log analysis, and completed security event documentation so endpoints stayed protected during active investigations",
            "Investigated suspicious file activity using Windows Event Logs and Wazuh, validated IOC matches with threat intelligence on VirusTotal, isolated affected workstations, and blocked malware from restarting at login",
            "Traced suspicious TCP/IP and DNS traffic from remote login alerts using IDS/IPS, Wireshark, and PCAP analysis, reviewed firewall and VPN logs to confirm the activity, and isolated affected hosts so the threat could not spread further across the network",
            "Ran authorized MITRE ATT&CK simulations on virtual machines, triaged alerts in Splunk and Wazuh with SPL and log analysis, and updated detection steps so repeat threats were caught before affecting end users",
            "Enriched IOCs with Python and PowerShell using threat intelligence from VirusTotal, documented each case in osTicket, and wrote plain-language incident escalation handoff notes so the next analyst could continue the investigation without losing context",
        ],
    },
    {
        "title": "Independent Microsoft Azure SOC Homelab",
        "dates": None,
        "bullets": [
            "Ran SOC operations in a Microsoft Azure environment with Microsoft Sentinel and Microsoft Defender for Endpoint on Entra ID workstations, performed alert triage through security monitoring and log analysis, and completed security event documentation for shift handoff so the next analyst could resume investigations without starting over",
            "Led a phishing investigation using Microsoft Sentinel and Microsoft Defender for Endpoint, reviewed HTTP/S email headers and endpoint logs through log analysis, and isolated the workstation so malware could not spread to other users",
            "Responded to a Microsoft 365 password spray using Entra ID logs and KQL in Microsoft Sentinel, revoked the active session and reset the compromised account through incident escalation, and blocked further access to email and file resources",
            "Tuned a KQL rule in Microsoft Sentinel using VPN logs after a firewall change, applied false positive reduction to close a noisy alert, and kept VPN security monitoring active so real remote access threats were still caught without alert fatigue",
            "Prioritized an alert queue in osTicket, triaged phishing and login alerts before lower-priority firewall noise, and closed tickets by risk through incident response so the highest-severity threats were handled first",
        ],
    },
]

WORK_HISTORY = [
    {
        "title": "Field/Technical Support at Hava Medical, Abuja, Nigeria",
        "dates": "February 2025 to Current",
        "bullets": [
            "Provided on-site field support for hospital medical equipment across multiple facilities, used troubleshooting and customer service to keep devices running, and coordinated with clinical and engineering teams so patient care was not delayed when equipment failed",
            "Explained technical specifications in plain language to non-technical hospital staff during product demonstrations, and supported new site adoption so facilities could use new equipment safely without waiting for another engineering visit",
        ],
    },
    {
        "title": "Freelance Copywriter, Remote",
        "dates": "March 2022 to September 2024",
        "bullets": [
            "Wrote research-based documentation for remote clients using Microsoft Office, structured findings into clear reports with consistent formatting, and delivered completed documents on deadline so managers could reference accurate written records without requesting follow-up edits",
        ],
    },
]


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    run.bold = bold
    run.italic = italic


def set_paragraph_format(paragraph, space_before=0, space_after=0, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if alignment is not None:
        fmt.alignment = alignment


def add_text_paragraph(doc, text, size=BODY_SIZE, bold=False, italic=False, space_before=0, space_after=0, alignment=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=space_before, space_after=space_after, alignment=alignment)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_section_heading(doc, text, space_before=8):
    return add_text_paragraph(doc, text, bold=True, space_before=space_before, space_after=2)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_format(p, space_after=0)
    if p.runs:
        p.clear()
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_job(doc, title, dates, bullets):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=2)
    if dates:
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        title_run = p.add_run(title)
        set_run_font(title_run, italic=True)
        p.add_run("\t")
        date_run = p.add_run(dates)
        set_run_font(date_run, italic=True)
    else:
        title_run = p.add_run(title)
        set_run_font(title_run, italic=True)

    for bullet in bullets:
        add_bullet(doc, bullet)


def configure_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = LINE_SPACING

    if "List Bullet" in doc.styles:
        style = doc.styles["List Bullet"]
        style.font.name = FONT
        style.font.size = BODY_SIZE
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = LINE_SPACING


def build():
    doc = Document()
    configure_document_styles(doc)

    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    add_text_paragraph(
        doc,
        CONTACT["name"],
        size=NAME_SIZE,
        bold=True,
        space_after=2,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text_paragraph(
        doc,
        f"{CONTACT['phone']} | {CONTACT['email']} | {CONTACT['links']}",
        size=CONTACT_SIZE,
        space_after=2,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text_paragraph(
        doc,
        CONTACT["location"],
        size=CONTACT_SIZE,
        space_after=6,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_section_heading(doc, "Education & Certificates", space_before=0)
    for line in EDUCATION:
        add_bullet(doc, line)

    add_section_heading(doc, "Projects")
    for project in PROJECTS:
        add_job(doc, project["title"], project["dates"], project["bullets"])

    add_section_heading(doc, "Work History", space_before=8)
    for job in WORK_HISTORY:
        add_job(doc, job["title"], job["dates"], job["bullets"])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Wrote {OUTPUT}")
    except PermissionError:
        doc.save(OUTPUT_FALLBACK)
        print(f"Wrote {OUTPUT_FALLBACK} (main file locked)")


if __name__ == "__main__":
    build()