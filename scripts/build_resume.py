#!/usr/bin/env python3
"""Build SOC Tier 1 resume as .docx (guide layout, standard bullet style)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "Praisel_Ekpenyong_Resume_SOC_Tier1_Analyst_v5.docx"

FONT = "Arial"
BLUE = RGBColor(0, 0, 255)


def set_font(run, size_pt: float, bold: bool = False, italic: bool = False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_center_line(doc, text: str, size: float, bold: bool = False, link: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size, bold=bold, color=BLUE if link else None)
    if link:
        run.font.underline = True


def add_section_header(doc, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_font(run, 10.5, bold=True)


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, 10.5)


def add_job_header(doc, left: str, right: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

    run_left = p.add_run(left)
    set_font(run_left, 10.5, italic=True)
    run_tab = p.add_run("\t")
    set_font(run_tab, 10.5)
    run_right = p.add_run(right)
    set_font(run_right, 10.5, italic=True)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # --- Section 1: Name, Contact, Location (center) ---
    add_center_line(doc, "Praisel Ekpenyong", 14, bold=True)

    contact = (
        "[Your Phone Number] | Ekpenyongpraisel@gmail.com | "
        "linkedin.com/in/praiselekpenyong | github.com/praisel-ekpenyong"
    )
    add_center_line(doc, contact, 12)

    # User should replace bracketed fields before submitting.
    add_center_line(doc, "[City], Alberta, Canada", 12)

    # --- Section 2: Education & Certificates ---
    add_section_header(doc, "Education & Certificates")
    add_bullet(doc, "CompTIA Security+ — CompTIA")
    add_bullet(doc, "Microsoft SC-200: Security Operations Analyst Associate — Microsoft")
    add_bullet(doc, "Google Cybersecurity Professional Certificate — Google")

    # --- Section 3: Projects (primary keyword density for SOC Tier 1) ---
    add_section_header(doc, "Projects")

    # Cloud lab first — tailored to MNP SOC Tier 1 (Sentinel, Defender, KQL, playbooks)
    add_job_header(
        doc,
        "SOC Analyst L1, Microsoft Cloud Security Operations Lab (Personal Project)",
        "2025 to Current",
    )
    add_bullet(
        doc,
        "Used Microsoft Sentinel and alert triage in the cloud incident queue to sort real threats "
        "from noise, so the team focused on problems that could hurt a client.",
    )
    add_bullet(
        doc,
        "Ran basic KQL in Microsoft Sentinel during investigations to pull process and login details, "
        "so I had proof before asking Tier 2 for help.",
    )
    add_bullet(
        doc,
        "Isolated a risky host with Microsoft Defender for Endpoint in the endpoint protection console "
        "within 13 minutes, so bad software could not spread to other computers.",
    )
    add_bullet(
        doc,
        "Followed playbooks and SOPs on the Tier 1 Security Operations ticket queue and escalated "
        "confirmed cases to Tier 2 with clear notes, so the next analyst and seniors had a clean handoff.",
    )
    add_bullet(
        doc,
        "Troubleshot a VPN login alert in Microsoft Sentinel by matching it to a firewall change ticket, "
        "so the night team did not chase a problem that was not real.",
    )

    # On-prem lab — 4 bullets
    add_job_header(
        doc,
        "SOC Analyst L1, On-Premises Security Operations Lab (Personal Project)",
        "2025 to Current",
    )
    add_bullet(
        doc,
        "Triaged SIEM alerts with Wazuh on the on-prem alert queue for Windows and Linux endpoints, "
        "so suspicious downloads and account changes were caught early.",
    )
    add_bullet(
        doc,
        "Used Sysmon and Windows Event Logs on finance team workstations to trace bad process chains, "
        "so we knew exactly which machine needed help.",
    )
    add_bullet(
        doc,
        "Reviewed network traffic with Splunk and Wireshark on the finance workstation segment during "
        "a suspected lateral movement case, so the IR team had clear evidence to act on.",
    )
    add_bullet(
        doc,
        "Tested detection rules with MITRE ATT&CK and Apache Caldera in the practice SOC lab, "
        "so we knew our alerts would fire during a real attack.",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()